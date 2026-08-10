"""Generate offer letter PDF/DOCX from resume extractions + HR offer fields."""
from __future__ import annotations

import html as html_module
import json
import re
from datetime import date
from io import BytesIO
from typing import Any

import pymupdf as fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _parse_extracted_data(raw: Any) -> dict:
    if isinstance(raw, dict):
        return raw
    if isinstance(raw, str) and raw.strip():
        try:
            parsed = json.loads(raw)
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}
    return {}


def merge_resume_extraction(extractions: list[dict]) -> dict[str, Any]:
    merged: dict[str, Any] = {}
    for row in extractions or []:
        if not isinstance(row, dict):
            continue
        data = _parse_extracted_data(row.get("extracted_data"))
        if not data:
            continue
        for key, val in data.items():
            if val is None or val == "" or val == []:
                continue
            if key not in merged or merged[key] in (None, "", []):
                merged[key] = val
    return merged


def prefill_from_resume(extractions: list[dict]) -> dict[str, Any]:
    data = merge_resume_extraction(extractions)
    name = (
        data.get("candidate_name")
        or data.get("full_name")
        or data.get("name")
        or ""
    )
    if isinstance(name, str):
        name = name.strip()
    job_title = data.get("current_title") or data.get("job_title") or ""
    if isinstance(job_title, str):
        job_title = job_title.strip()
    cnic = (
        data.get("cnic")
        or data.get("cnic_passport")
        or data.get("national_id")
        or data.get("nic")
        or ""
    )
    if isinstance(cnic, str):
        cnic = cnic.strip()
    return {
        "candidate_name": name or None,
        "job_title": job_title or None,
        "email": data.get("email"),
        "phone": data.get("phone"),
        "location": data.get("location"),
        "cnic": cnic or None,
        "resume_summary": data.get("evaluation_summary"),
        "source_fields_used": list(data.keys())[:20],
    }


def _fmt_money(amount: Any, currency: str) -> str:
    try:
        num = float(amount)
        cur = (currency or "").strip().upper()
        if cur == "PKR":
            return f"{cur} {num:,.0f}"
        return f"{cur} {num:,.2f}".strip() if cur else f"{num:,.2f}"
    except (TypeError, ValueError):
        return str(amount or "")


def _parse_iso_date(s: str | None) -> date | None:
    if not s or not str(s).strip():
        return None
    try:
        return date.fromisoformat(str(s).strip()[:10])
    except ValueError:
        return None


def validate_offer_dates(offer: dict[str, Any]) -> list[str]:
    """Return human-readable validation errors (Pakistan one-page flow)."""
    errors: list[str] = []
    letter_d = _parse_iso_date(offer.get("letter_date"))
    joining_d = _parse_iso_date(offer.get("joining_date"))
    valid_d = _parse_iso_date(offer.get("offer_valid_until"))
    today = date.today()
    if letter_d and letter_d > today:
        errors.append("Letter date cannot be in the future.")
    if letter_d and joining_d and joining_d < letter_d:
        errors.append("Joining date must be on or after the letter date.")
    if letter_d and valid_d and valid_d < letter_d:
        errors.append("Offer valid-until date must be on or after the letter date.")
    return errors


def _normalize_offer(offer: dict[str, Any], resume_hint: dict[str, Any] | None) -> dict[str, str]:
    resume_hint = resume_hint or {}
    pay_freq = (offer.get("pay_frequency") or "Monthly").strip()
    currency = (offer.get("currency") or "PKR").strip()
    salary_raw = offer.get("offered_salary")
    salary_line = ""
    if salary_raw not in (None, "", 0):
        salary_line = f"{_fmt_money(salary_raw, currency)} per {pay_freq.lower()}"
    else:
        salary_line = "As mutually agreed (to be confirmed in writing before joining)"

    return {
        "candidate": (offer.get("candidate_name") or resume_hint.get("candidate_name") or "Candidate").strip(),
        "company": (offer.get("company_name") or "Company").strip(),
        "company_address": (offer.get("company_address") or "").strip(),
        "job_title": (offer.get("job_title") or resume_hint.get("job_title") or "Role").strip(),
        "department": (offer.get("department") or "").strip(),
        "location": (offer.get("work_location") or resume_hint.get("location") or "").strip(),
        "joining": (offer.get("joining_date") or "").strip(),
        "valid_until": (offer.get("offer_valid_until") or "").strip(),
        "probation": (offer.get("probation_period") or "3 months").strip(),
        "notice_period": (offer.get("notice_period") or "30 days").strip(),
        "pay_freq": pay_freq,
        "currency": currency,
        "salary_line": salary_line,
        "cnic": (offer.get("cnic") or resume_hint.get("cnic") or "").strip(),
        "signatory_name": (offer.get("signatory_name") or "Human Resources").strip(),
        "signatory_title": (offer.get("signatory_title") or "Authorized Signatory").strip(),
        "notes": (offer.get("additional_notes") or "").strip(),
        "letter_date": (offer.get("letter_date") or date.today().isoformat()).strip(),
        "email": str(resume_hint.get("email") or offer.get("email") or "").strip(),
        "phone": str(resume_hint.get("phone") or offer.get("phone") or "").strip(),
        "include_background": bool(offer.get("include_background")),
        "summary": str(resume_hint.get("resume_summary") or "").strip()[:800],
    }


def _term_line(label: str, value: str, esc) -> str:
    if not value or value == "—":
        return ""
    return (
        f'<p style="margin:0 0 5px;font-size:9.5pt;">'
        f'<span style="font-weight:700;">{esc(label)}:</span> {value}</p>'
    )


def build_offer_letter_html(offer: dict[str, Any], resume_hint: dict[str, Any] | None = None) -> str:
    """Pakistan-style one-page offer letter (no HTML tables — PyMuPDF-safe layout)."""
    f = _normalize_offer(offer, resume_hint)
    esc = html_module.escape

    dept_part = f" ({esc(f['department'])})" if f["department"] else ""
    loc_part = f" at {esc(f['location'])}" if f["location"] else ""

    contact_bits = []
    if f["email"]:
        contact_bits.append(f"Email: {esc(f['email'])}")
    if f["phone"]:
        contact_bits.append(f"Phone: {esc(f['phone'])}")
    if f["cnic"]:
        contact_bits.append(f"CNIC: {esc(f['cnic'])}")
    contact_line = " · ".join(contact_bits)

    addr_block = ""
    if f["company_address"]:
        addr_block = f'<p style="margin:2px 0 0;font-size:8.5pt;color:#475569;">{esc(f["company_address"])}</p>'

    terms = "".join(
        filter(
            None,
            [
                _term_line("Position", esc(f["job_title"]), esc),
                _term_line("Department", esc(f["department"]) if f["department"] else "", esc),
                _term_line("Work location", esc(f["location"]) if f["location"] else "", esc),
                _term_line("Compensation (gross)", esc(f["salary_line"]), esc),
                _term_line("Expected joining date", esc(f["joining"]) if f["joining"] else "To be confirmed", esc),
                _term_line("Probation period", esc(f["probation"]), esc),
                _term_line("Notice period", esc(f["notice_period"]), esc),
                _term_line("Offer valid until", esc(f["valid_until"]) if f["valid_until"] else "—", esc),
            ],
        )
    )

    notes_block = ""
    if f["notes"]:
        notes_block = f'<p style="margin:8px 0 0;font-size:9pt;text-align:justify;"><b>Note:</b> {esc(f["notes"])}</p>'

    background_block = ""
    if f["include_background"] and f["summary"]:
        background_block = (
            f'<p style="margin:8px 0 0;font-size:9pt;text-align:justify;">'
            f"<b>Background:</b> {esc(f['summary'])}</p>"
        )

    valid_accept = esc(f["valid_until"]) if f["valid_until"] else "the date stated above"

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"/></head>
<body style="font-family:Helvetica,Arial,sans-serif;font-size:9.5pt;line-height:1.38;color:#0f172a;margin:0;">
  <div style="text-align:center;margin-bottom:10px;border-bottom:2px solid #1e3a5f;padding-bottom:8px;">
    <div style="font-size:14pt;font-weight:bold;color:#1e3a5f;">{esc(f['company'])}</div>
    {addr_block}
    <div style="font-size:10pt;font-weight:600;margin-top:6px;letter-spacing:0.02em;">OFFER OF EMPLOYMENT</div>
  </div>
  <p style="margin:0 0 6px;"><b>Date:</b> {esc(f['letter_date'])}</p>
  <p style="margin:0 0 4px;"><b>To:</b> {esc(f['candidate'])}</p>
  {f'<p style="margin:0 0 8px;font-size:9pt;color:#334155;">{contact_line}</p>' if contact_line else '<p style="margin:0 0 8px;"></p>'}
  <p style="margin:0 0 6px;">Dear {esc(f['candidate'])},</p>
  <p style="margin:0 0 8px;text-align:justify;">
    We are pleased to offer you employment with <b>{esc(f['company'])}</b> as <b>{esc(f['job_title'])}</b>{dept_part}{loc_part}.
    This letter summarizes the principal terms of your offer in accordance with our company policies and applicable laws of Pakistan.
  </p>
  <p style="margin:0 0 4px;font-weight:700;font-size:10pt;">Terms of offer</p>
  {terms}
  <p style="margin:8px 0 0;text-align:justify;font-size:9pt;">
    You will be entitled to leave and holidays as per company policy and applicable labour laws. Working hours and benefits
    will be as communicated during onboarding. This offer is subject to verification of your credentials, medical fitness where
    required, and submission of documents (including CNIC and educational certificates). You will maintain confidentiality of
    company information during and after employment.
  </p>
  {notes_block}
  {background_block}
  <p style="margin:10px 0 8px;text-align:justify;font-size:9pt;">
    Please sign below to accept this offer and return a copy by <b>{valid_accept}</b>. We look forward to you joining our team.
  </p>
  <p style="margin:0 0 20px;">Yours sincerely,</p>
  <p style="margin:0 0 2px;font-weight:700;">{esc(f['signatory_name'])}</p>
  <p style="margin:0 0 14px;font-size:9pt;">{esc(f['signatory_title'])} · {esc(f['company'])}</p>
  <div style="margin-top:12px;padding-top:8px;border-top:1px solid #cbd5e1;">
    <p style="margin:0 0 28px;font-size:9pt;font-weight:700;">Candidate acceptance</p>
    <p style="margin:0;font-size:9pt;">I, <b>{esc(f['candidate'])}</b>, accept the above terms of employment.</p>
    <p style="margin:12px 0 0;font-size:9pt;">Signature: _________________________ &nbsp;&nbsp; Date: _______________</p>
  </div>
</body></html>"""


def generate_offer_letter_pdf(offer: dict[str, Any], resume_hint: dict[str, Any] | None = None) -> bytes:
    errors = validate_offer_dates(offer)
    if errors:
        raise ValueError("; ".join(errors))

    html = build_offer_letter_html(offer, resume_hint)
    mediabox = fitz.paper_rect("a4")
    where = mediabox + (42, 42, -42, -42)
    buf = BytesIO()
    writer = fitz.DocumentWriter(buf)
    story = fitz.Story(
        html=html,
        user_css="body { font-family: Helvetica, Arial, sans-serif; font-size: 9.5pt; }",
    )
    while True:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
        if not more:
            break
    writer.close()
    return buf.getvalue()


def generate_offer_letter_docx(offer: dict[str, Any], resume_hint: dict[str, Any] | None = None) -> bytes:
    f = _normalize_offer(offer, resume_hint)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f["company"])
    run.bold = True
    run.font.size = Pt(14)
    sub = doc.add_paragraph("OFFER OF EMPLOYMENT")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if f["company_address"]:
        p = doc.add_paragraph(f["company_address"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {f['letter_date']}")
    doc.add_paragraph(f"To: {f['candidate']}")
    if f["cnic"]:
        doc.add_paragraph(f"CNIC: {f['cnic']}")
    doc.add_paragraph(f"Dear {f['candidate']},")
    doc.add_paragraph(
        f"We are pleased to offer you employment with {f['company']} as {f['job_title']}."
    )
    doc.add_paragraph("Terms of offer", style="Heading 3")
    for label, val in [
        ("Position", f["job_title"]),
        ("Department", f["department"] or "—"),
        ("Work location", f["location"] or "—"),
        ("Compensation", f["salary_line"]),
        ("Joining date", f["joining"] or "TBC"),
        ("Probation", f["probation"]),
        ("Notice period", f["notice_period"]),
        ("Valid until", f["valid_until"] or "—"),
    ]:
        doc.add_paragraph(f"{label}: {val}")
    if f["notes"]:
        doc.add_paragraph(f"Note: {f['notes']}")
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph(f"{f['signatory_name']}\n{f['signatory_title']}")
    doc.add_paragraph("Candidate acceptance: _________________________  Date: _______________")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def safe_filename(candidate: str, ext: str = "pdf") -> str:
    base = re.sub(r"[^\w\s-]", "", candidate or "Candidate").strip().replace(" ", "_")
    if not base:
        base = "Candidate"
    ext = ext.lstrip(".") or "pdf"
    return f"Offer_Letter_{base[:48]}.{ext}"
